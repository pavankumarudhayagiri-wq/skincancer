"""
Generate B.Tech project report DOCX for Melanoma Detection App.
Structure mirrors C-G1 ProjectDoc1.pdf; content reflects only tools/models in this repo.
Run: pip install python-docx matplotlib
     python generate_melanoma_report_docx.py
"""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

OUT_DIR = Path(__file__).resolve().parent
FIG_DIR = OUT_DIR / "report_assets"
DOCX_PATH = OUT_DIR / "Melanoma_Detection_Project_Report.docx"

# Models shipped / selectable in app (from melanoma_ml.py)
MODEL_NAMES = (
    "CNN",
    "VGG16",
    "ResNet50",
    "EfficientNetB4",
    "InceptionResNetV2",
)

# requirements.txt only (plus Python stdlib where relevant)
PIP_STACK = [
    ("streamlit", "Web interface and page navigation"),
    ("numpy", "Numerical arrays and image tensors"),
    ("matplotlib", "Static plots and figures"),
    ("pandas", "Tabular metrics and dataset CSV handling"),
    ("seaborn", "Statistical styling for charts"),
    ("plotly", "Interactive visualizations"),
    ("pyyaml", "YAML configuration"),
    ("pillow", "Image loading and drawing"),
    ("scikit-learn", "ROC curves, AUC, label utilities"),
    ("gdown", "Downloading shared model assets when needed"),
    ("tensorflow-cpu (>=2.10)", "Keras models, training artifacts, inference"),
    ("opencv-python", "Image processing when available"),
    ("pytesseract", "Optional OCR-related helpers when Tesseract is installed"),
    ("scipy", "Scientific helpers such as image filters"),
]


def set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_heading(doc: Document, text: str, caps: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    t = text.upper() if caps else text
    run = p.add_run(t)
    set_run_font(run, 14, bold=True)


def add_body(doc: Document, text: str, justify: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 12, bold=False)


def add_title_block(doc: Document) -> None:
    for line in [
        "A PROJECT REPORT",
        "ON",
        "MELANOMA DETECTION WEB APPLICATION",
        "",
        "A project report submitted in partial fulfillment of the requirements for the award of the degree of",
        "BACHELOR OF TECHNOLOGY",
        "IN",
        "COMPUTER SCIENCE & ENGINEERING",
        "",
        "Submitted By",
        "S. RUCHITHA 212P1A05D3",
        "S. SARANYA 212P1A05D5",
        "S. MUBARAK 212P1A05C6",
        "S. KALYANI 212P1A05D6",
        "V. THANMAI 212P1A05E6",
        "",
        "Under The Esteemed Guidance Of",
        "Mrs. S. Himaja, M.Tech",
        "Asst.Professor",
        "",
        "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING",
        "AN ISO 9001:2015 CERTIFIED INSTITUTION",
        "CHAITANYABHARATHI INSTITUTE OF TECHNOLOGY (AUTONOMOUS)",
        "(Sponsored by Bharathi Educational Society)",
        "(Affiliated to J.N.T.U.A., Anantapuramu, Approved by AICTE, New Delhi)",
        "Recognized by UGC Under the Sections 2(f)&12(B) of UGC Act, 1956",
        "(Accredited by NAAC & NBA)",
        "Vidyanagar, Proddatur-516360, Y.S.R.(Dist.), A.P",
        "2021-2025",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(line)
        set_run_font(run, 16 if line == "MELANOMA DETECTION WEB APPLICATION" else 12, bold=line == "MELANOMA DETECTION WEB APPLICATION")
    doc.add_page_break()


def add_certificate_declaration_ack(doc: Document) -> None:
    add_heading(doc, "CERTIFICATE", caps=True)
    add_body(
        doc,
        'This is to certify that the project work entitled "MELANOMA DETECTION WEB APPLICATION" is a '
        "Bonafide work of S. RUCHITHA (212P1A05D3), S. SARANYA REDDY (212P1A05D5), S. MUBARAK (212P1A05C6), "
        "S. KALYANI (212P1A05D6), V. THANMAI (212P1A05E6) submitted to Chaitanya Bharathi Institute of Technology, "
        "Proddatur in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in "
        "COMPUTER SCIENCE AND ENGINEERING. The work reported herein does not form part of any other thesis on which "
        "a degree has been awarded earlier. This is to further certify that they have worked for a period of one "
        "semester for preparing their work under our supervision and guidance.",
    )
    add_body(doc, "INTERNAL GUIDE\t\tHEAD OF THE DEPARTMENT")
    add_body(doc, "Mrs. S. Himaja, M.Tech\tDr. Y. Dasaratha Rami Reddy, M.Tech, Ph.D")
    add_body(doc, "Asst.Professor\t\tProfessor")
    add_body(doc, "PROJECT CO-ORDINATOR")
    add_body(doc, "Dr. P. Narasimhaiah, M.Tech, Ph.D,")
    add_body(doc, "Associate Professor")
    add_body(doc, "INTERNAL EXAMINER\t\tEXTERNAL EXAMINER")
    doc.add_page_break()

    add_heading(doc, "DECLARATION BY THE CANDIDATES", caps=True)
    add_body(
        doc,
        "We are S. Ruchitha, S. Saranya Reddy, S. Mubarak, S. Kalyani, V. Thanmai with respective Roll No: "
        "(212P1A05D3), (212P1A05D5), (212P1A05C6), (212P1A05D6), (212P1A05E6) hereby declare that the Project Report "
        'entitled "MELANOMA DETECTION WEB APPLICATION" under the guidance of Mrs. S. Himaja, Asst.Professor, '
        "Department of CSE is submitted in partial fulfillment of the requirements for the award of the degree of "
        "Bachelor of Technology in Computer Science & Engineering.\n"
        "This is a record of bonafide work carried out by us and the results embodied in this Project Report have "
        "not been reproduced or copied from any source. The results embodied in this Project Report have not been "
        "submitted to any other University or Institute for the Award of any other Degree or Diploma.",
    )
    add_body(doc, "S. RUCHITHA 212P1A05D3\nS. SARANYA REDDY 212P1A05D5\nS. MUBARAK 212P1A05C6\nS. KALYANI 212P1A05D6\nV. THANMAI 212P1A05E6")
    add_body(doc, "Dept. of Computer Science & Engineering\nChaitanya Bharathi Institute of Technology\nVidyanagar, Proddatur, Y.S.R.(Dist.)")
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT", caps=True)
    add_body(
        doc,
        "An endeavour over a long period can be successful only with the advice and support of many well-wishers. "
        "We take this opportunity to express our gratitude and appreciation to all of them.\n"
        "We are extremely thankful to our beloved Chairman, Dr. V. Jayachandra Reddy, who took keen interest and "
        "encouraged us in every effort throughout this course.\n"
        "We would like to thank our esteemed Director (Admin), Dr. G. Sreenivasula Reddy, M.Tech., Ph.D., who have "
        "truly enriched our understanding and inspired us.\n"
        "We owe our gratitude to our Principal Dr. S. Sruthi, M.Tech., Ph.D. for permitting us to use the facilities "
        "available to accomplish the project successfully.\n"
        "We express our heartfelt thanks to Dr. Y. Dasaratha Rami Reddy, M.Tech., Ph.D., Head of Dept. of CSE for his "
        "kind attention and valuable guidance to us throughout this course.\n"
        "We also express our deep sense of gratitude towards Mrs. S. Himaja, M.Tech., Asst.Professor, Dept. of CSE, "
        "for her support and guidance in completing our project.\n"
        "We express our profound gratitude to our project coordinator Dr. P. Narasimhaiah, M.Tech, Ph.D., for his "
        "valuable support and guidance in completing the project successfully.\n"
        "We also thank all the teaching & non-teaching staff of the Dept. of CSE for their support throughout our "
        "B.Tech course.\n"
        "We express our heartfelt thanks to our parents for their valuable support and encouragement in completion of "
        "our course. Also, we express our heartfelt regards to our friends for being supportive in completion of the project.",
    )
    doc.add_page_break()


def build_flowchart_png(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def box(x, y, w, h, text):
        r = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.15",
            linewidth=1.2, edgecolor="#1a365d", facecolor="#ebf8ff",
        )
        ax.add_patch(r)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9, wrap=True)

    def arrow(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", color="#2c5282", lw=1.2))

    box(3.5, 9.0, 3.0, 0.7, "Open web app\n(Streamlit)")
    arrow(5, 9.0, 5, 8.45)
    box(3.3, 7.6, 3.4, 0.9, "Login / Register\n(SQLite user store)")
    arrow(5, 7.6, 5, 7.0)
    box(3.2, 6.1, 3.6, 0.9, "Authenticated dashboard\n(sidebar pages)")
    arrow(5, 6.1, 5, 5.5)
    box(2.0, 4.7, 6.0, 0.8, "Optional learning path:\nIntro, metrics, PH2 / ISIC2016 charts")
    arrow(5, 4.7, 5, 4.15)
    box(2.5, 3.3, 5.0, 0.85, "Detection: upload image → choose skin or dermoscopy")
    arrow(5, 3.3, 5, 2.75)
    box(2.2, 1.9, 5.6, 0.85, "Select model (" + ", ".join(MODEL_NAMES[:3]) + ", …) → TensorFlow inference")
    arrow(5, 1.9, 5, 1.35)
    box(3.0, 0.5, 4.0, 0.75, "Show probability, resources, FAQ, feedback → Logout")

    ax.set_title("Figure 3.5 — Project workflow (Melanoma Detection Web Application)", fontsize=11, fontweight="bold")
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def build_architecture_png(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis("off")

    layers = [
        (0.5, 6.5, 9, 1.0, "Client: Web browser"),
        (0.5, 5.0, 9, 1.0, "Presentation: Streamlit (Python) — forms, charts, file upload"),
        (0.5, 3.5, 9, 1.0, "Application: melanoma_ml.py — auth, routing, metrics, visualization, inference orchestration"),
        (0.5, 2.0, 4.2, 1.0, "Data: SQLite (users.db) — registered users"),
        (5.3, 2.0, 4.2, 1.0, "Artifacts: .keras / .h5 weights, CSV ground-truth (ISIC2016), PH2-related data files"),
        (0.5, 0.5, 9, 1.0, "ML runtime: TensorFlow Keras — " + ", ".join(MODEL_NAMES)),
    ]
    colors = ["#e6fffa", "#ebf8ff", "#faf5ff", "#fefcbf", "#feebc8", "#e9d8fd"]
    for (x, y, w, h, txt), c in zip(layers, colors):
        r = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.1",
            linewidth=1.2, edgecolor="#2d3748", facecolor=c,
        )
        ax.add_patch(r)
        ax.text(x + w / 2, y + h / 2, txt, ha="center", va="center", fontsize=8.5)

    ax.annotate(
        "",
        xy=(5, 6.5), xytext=(5, 6.0),
        arrowprops=dict(arrowstyle="->", color="#2c5282", lw=1.5),
    )
    ax.annotate(
        "",
        xy=(5, 5.0), xytext=(5, 4.5),
        arrowprops=dict(arrowstyle="->", color="#2c5282", lw=1.5),
    )
    ax.annotate(
        "",
        xy=(2.6, 3.5), xytext=(2.6, 3.0),
        arrowprops=dict(arrowstyle="->", color="#2c5282", lw=1.2),
    )
    ax.annotate(
        "",
        xy=(7.4, 3.5), xytext=(7.4, 3.0),
        arrowprops=dict(arrowstyle="->", color="#2c5282", lw=1.2),
    )
    ax.annotate(
        "",
        xy=(5, 2.0), xytext=(5, 1.5),
        arrowprops=dict(arrowstyle="->", color="#2c5282", lw=1.2),
    )

    ax.set_title("Figure 3.6 — System architecture (logical layers)", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def add_index_and_lists(doc: Document) -> None:
    add_heading(doc, "INDEX", caps=True)
    rows = [
        ("ABSTRACT", ""),
        ("1. INTRODUCTION", ""),
        ("1.1 Problem Statement", ""),
        ("1.2 Motivation", ""),
        ("1.3 Objective", ""),
        ("1.4 Scope", ""),
        ("1.5 Project Introduction", ""),
        ("2. LITERATURE SURVEY", ""),
        ("2.1 Related Work", ""),
        ("3. SYSTEM ANALYSIS", ""),
        ("3.1 Existing System", ""),
        ("3.2 Disadvantages", ""),
        ("3.3 Proposed System", ""),
        ("3.4 Advantages", ""),
        ("3.5 Project flow", ""),
        ("3.6 Architecture", ""),
        ("4. REQUIREMENT ANALYSIS", ""),
        ("4.1 Hardware Requirements", ""),
        ("4.2 Software Requirements", ""),
        ("5. SYSTEM DESIGN", ""),
        ("5.1 Introduction of input design", ""),
        ("5.2 UML Diagrams (Use case, Class, Sequence, Activity,)", ""),
        ("6. IMPLEMENTATION AND RESULTS", ""),
        ("6.1 Modules", ""),
        ("6.2 Results", ""),
        ("7. TECHNOLOGIES USED", ""),
        ("8. SYSTEM STUDY AND TESTING", ""),
        ("8.1 Feasibility study", ""),
        ("8.2 Types of Testing", ""),
        ("8.3 Test Cases", ""),
        ("9. CONCLUSION", ""),
        ("10. FUTURE ENHANCEMENT", ""),
        ("11. REFERENCES", ""),
        ("12. APPENDIX", ""),
        ("13. BIO DATA", ""),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "CONTENT"
    hdr[1].text = "PAGE NO"
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    doc.add_page_break()

    add_heading(doc, "LIST OF FIGURES", caps=True)
    fig_rows = [
        ("1", "3.5", "Workflow Diagram", "See Figure 3.5"),
        ("2", "3.6", "Architecture", "See Figure 3.6"),
        ("3", "5.2.1", "Use case diagram", "Prepared separately / department template"),
        ("4", "5.2.2", "Class diagram", "Prepared separately / department template"),
        ("5", "5.2.3", "Sequence diagram", "Prepared separately / department template"),
        ("6", "5.2.6", "Activity diagram", "Prepared separately / department template"),
    ]
    t2 = doc.add_table(rows=1 + len(fig_rows), cols=4)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    for j, h in enumerate(["S.NO", "FIGURE NO.", "FIGURE NAME", "PAGE NO."]):
        h2[j].text = h
    for i, row in enumerate(fig_rows, start=1):
        for j, val in enumerate(row):
            t2.rows[i].cells[j].text = val
    doc.add_page_break()

    add_heading(doc, "LIST OF OUTPUT SCREENS", caps=True)
    screens = [
        ("1", "6.2.1", "Authentication / entry (Login & Register tabs)", ""),
        ("2", "6.2.2", "Introduction page", ""),
        ("3", "6.2.3", "Model performance — metrics tables", ""),
        ("4", "6.2.4", "Visualizations — PH2 / ISIC2016 distributions", ""),
        ("5", "6.2.5", "Melanoma detection — skin image upload & model choice", ""),
        ("6", "6.2.6", "Melanoma detection — dermoscopy upload & model choice", ""),
        ("7", "6.2.7", "Prediction output with probability", ""),
        ("8", "6.2.8", "Model summary / architecture text download (where enabled)", ""),
        ("9", "6.2.9", "Educational resources", ""),
        ("10", "6.2.10", "FAQs", ""),
        ("11", "6.2.11", "Feedback and contact", ""),
        ("12", "6.2.12", "Session logout / re-login", ""),
        ("13", "6.2.13", "Forgot-password / recovery flow (if configured)", ""),
        ("14", "6.2.14", "Error state — missing model weights", ""),
        ("15", "6.2.15", "Error state — invalid upload type", ""),
        ("16", "6.2.16", "ROC / AUC style plots (evaluation section)", ""),
        ("17", "6.2.17", "Sidebar navigation overview", ""),
        ("18", "6.2.18", "Disclaimer / non-diagnostic notice", ""),
        ("19", "6.2.19", "Wide layout dashboard shell", ""),
    ]
    t3 = doc.add_table(rows=1 + len(screens), cols=4)
    t3.style = "Table Grid"
    h3 = t3.rows[0].cells
    for j, h in enumerate(["S.NO", "Figure No", "Figure Name", "PageNo"]):
        h3[j].text = h
    for i, row in enumerate(screens, start=1):
        for j, val in enumerate(row):
            t3.rows[i].cells[j].text = val
    doc.add_page_break()


def add_main_chapters(doc: Document, flow_png: Path, arch_png: Path) -> None:
    add_heading(doc, "ABSTRACT", caps=True)
    add_body(
        doc,
        "Melanoma is a dangerous form of skin cancer where earlier attention often improves outcomes. Many learners still "
        "study the topic through scattered articles and static slides, with little connection to how modern deep learning "
        "models behave on real skin and dermoscopy images. This project implements a browser-based application that keeps "
        "the workflow simple: users register and sign in, explore model performance and dataset charts, then upload an image "
        "and obtain a probability-style prediction from one of several Keras models. The stack is grounded in the actual "
        f"repository: Streamlit for the interface, Python for orchestration, SQLite for user records, and TensorFlow for "
        f"inference. The application supports the following model options: {', '.join(MODEL_NAMES)}. Training-related "
        "context uses PH2 and ISIC2016 materials where those files are present in the project. The tool is educational "
        "and includes clear messaging that it does not replace a qualified clinician.",
    )
    doc.add_paragraph()

    add_heading(doc, "1. INTRODUCTION", caps=True)
    add_heading(doc, "1.1 PROBLEM STATEMENT", caps=True)
    add_body(
        doc,
        "Students and interested users need a single place to see how melanoma classification models are evaluated, how "
        "datasets differ, and what happens when a new lesion photo is passed through a trained network. Without such a "
        "workspace, understanding stays theoretical and the gap between lecture slides and deployable software remains wide.",
    )
    add_heading(doc, "1.2 MOTIVATION", caps=True)
    add_body(
        doc,
        "We chose a web application because it is the most natural way to share an interactive demo. Streamlit allows the "
        "team to focus on reliable inference and honest metrics instead of rebuilding low-level web plumbing. The motivation "
        "is to make responsible, explainable use of TensorFlow models tangible in a final-year project setting.",
    )
    add_heading(doc, "1.3 OBJECTIVE OF THE PROJECT", caps=True)
    add_body(
        doc,
        "The objectives are: (1) provide secure registration and login backed by SQLite; (2) expose each implemented model "
        f"({', '.join(MODEL_NAMES)}) for both skin and dermoscopy tracks as defined in code; (3) visualize dataset and "
        "performance information using pandas, seaborn, and plotly; (4) keep the interface approachable for non-experts while "
        "showing medical disclaimers; (5) allow structured feedback so the application can evolve.",
    )
    add_heading(doc, "1.4 SCOPE", caps=True)
    add_body(
        doc,
        "The scope covers an educational web workflow on a single Python codebase. It does not claim regulatory approval. "
        "Future work may add hospital-grade audit logs or mobile clients, but those are outside the current implementation.",
    )
    add_heading(doc, "1.5 PROJECT INTRODUCTION", caps=True)
    add_body(
        doc,
        "The Melanoma Detection Web Application is implemented mainly in melanoma_ml.py with a thin Streamlit entry file. "
        "After authentication, users move through sidebar sections for introduction content, metric review, interactive plots, "
        "and the detection page. Weights are loaded from the .keras and .weights.h5 filenames declared in the source code, "
        "and optional gdown support helps when models are fetched from shared storage during setup.",
    )

    add_heading(doc, "2. LITERATURE SURVEY", caps=True)
    add_heading(doc, "2.1 RELATED WORK", caps=True)
    add_body(
        doc,
        "[1] Esteva, A., et al. (2017). Dermatologist-level classification of skin cancer with deep neural networks. Nature. "
        "This work motivates CNN-based screening aids.\n"
        "[2] Codella, N., et al. (2018). Skin lesion analysis toward melanoma detection (ISIC challenge lineage). "
        "Relevant to ISIC2016-style evaluation used in our visualization module.\n"
        "[3] Tschandl, P., et al. (2018). HAM10000 dataset paper, Scientific Data. Useful background on dermatoscopic data "
        "organization (conceptual reference for readers).\n"
        "[4] ISIC Archive public documentation — supports the ISIC2016 ground-truth CSV used in the project when present.",
    )

    add_heading(doc, "3. SYSTEM ANALYSIS", caps=True)
    add_heading(doc, "3.1 EXISTING SYSTEM", caps=True)
    add_body(
        doc,
        "Typical learning paths rely on notebooks or static PDFs. They rarely combine authenticated access, multi-model "
        "inference, and dataset charts in one Streamlit experience tied to concrete weight files.",
    )
    add_heading(doc, "3.2 DISADVANTAGES", caps=True)
    add_body(
        doc,
        "• Fragmented materials reduce hands-on practice.\n"
        "• Models shown only in slides do not teach preprocessing or failure modes.\n"
        "• Public tools without disclaimers may be misunderstood as diagnostic devices.",
    )
    add_heading(doc, "3.3 PROPOSED SYSTEM", caps=True)
    add_body(
        doc,
        "The proposed system is a single Python Streamlit application with SQLite authentication, TensorFlow inference, "
        "and libraries from requirements.txt for plotting and tables.",
    )
    add_heading(doc, "3.4 ADVANTAGES", caps=True)
    add_body(
        doc,
        "3.4.1 One coherent workflow from login to prediction.\n"
        "3.4.2 Multiple architectures can be compared on the same upload for learning.\n"
        "3.4.3 Lightweight deployment: users only need a browser once the server is running.\n"
        "3.4.4 Charts connect raw CSV ground truth to what the user sees on screen.",
    )
    add_heading(doc, "3.5 PROJECT FLOW", caps=True)
    add_body(doc, "The workflow figure below was generated for this report to match section 3.5 of the reference template.")
    doc.add_picture(str(flow_png), width=Inches(6.0))
    add_heading(doc, "3.6 ARCHITECTURE", caps=True)
    add_body(doc, "The architecture figure summarizes layers actually used in this repository (browser, Streamlit, Python logic, SQLite, files, TensorFlow).")
    doc.add_picture(str(arch_png), width=Inches(6.2))

    add_heading(doc, "4. REQUIREMENT ANALYSIS", caps=True)
    add_heading(doc, "4.1 HARDWARE REQUIREMENTS", caps=True)
    add_body(
        doc,
        "• Processor: Intel Core i5 or equivalent (CPU inference is configured in code; GPU is optional).\n"
        "• RAM: 8 GB minimum; 16 GB recommended for comfortable model loading.\n"
        "• Storage: sufficient space for Python environment plus .keras/.h5 weights.\n"
        "• Network: only if optional downloads (gdown) or email feedback is enabled.",
    )
    add_heading(doc, "4.2 SOFTWARE REQUIREMENTS", caps=True)
    add_body(
        doc,
        "• Operating System: Windows 10/11, Linux, or macOS.\n"
        "• Python 3.x with packages exactly as listed in requirements.txt (see Section 7).\n"
        "• Modern web browser.\n"
        "• Optional: Tesseract OCR binary if pytesseract paths are used on Windows.",
    )

    add_heading(doc, "5. SYSTEM DESIGN", caps=True)
    add_heading(doc, "5.1 INTRODUCTION OF INPUT DESIGN", caps=True)
    add_body(
        doc,
        "Registration captures username, email, and password with validation helpers. Login accepts username or email. "
        "Detection inputs are image uploads plus explicit choices for skin versus dermoscopy and for which of the five "
        "models should run. Validation reduces empty submissions and mismatched password flows.",
    )
    add_heading(doc, "5.2 UML DIAGRAMS", caps=True)
    add_body(
        doc,
        "5.2.1 Use case diagram: actors include Registered User and System processes for authenticate, view metrics, "
        "visualize datasets, run detection, and submit feedback.\n"
        "5.2.2 Class diagram: group entities such as User, SessionState, ModelArtifact, PredictionResult, and Feedback.\n"
        "5.2.3 Sequence diagram: login request, credential check against SQLite, page render, upload, preprocess, predict, display.\n"
        "5.2.4 Activity diagram: decision nodes for authenticated, valid file type, weights present, inference success.\n"
        "(Drawings can follow the department’s UML template; figures 5.2.1–5.2.6 remain listed in the List of Figures.)",
    )

    add_heading(doc, "6. IMPLEMENTATION AND RESULTS", caps=True)
    add_heading(doc, "6.1 MODULES", caps=True)
    add_body(
        doc,
        "6.1 USER MODULE\n"
        "• Register and login via SQLite-backed accounts.\n"
        "• Navigate Streamlit sidebar sections.\n"
        "• View model metrics and educational text.\n"
        "• Upload skin or dermoscopy images; pick CNN, VGG16, ResNet50, EfficientNetB4, or InceptionResNetV2.\n"
        "• Read FAQs, external resources, and submit feedback.\n"
        "• Logout / session reset.\n\n"
        "6.1.1 MODEL AND ANALYTICS MODULE\n"
        "• Loads Keras models or weight files named in melanoma_ml.py.\n"
        "• Applies backbone-specific preprocessing functions from tensorflow.keras.applications.\n"
        "• Builds plots with matplotlib, seaborn, and plotly; ROC utilities use scikit-learn.\n"
        "• Reads ISIC2016 CSV and PH2-related data when files exist in the working directory.",
    )
    add_heading(doc, "6.2 RESULTS", caps=True)
    add_body(
        doc,
        "Attach screenshots for screens listed in the List of Output Screens. Each figure should show the Streamlit caption "
        "and a short note describing the user action (for example, successful prediction on a dermoscopy sample). "
        "Placeholder lines are intentionally left here for pasting images in Microsoft Word if you extend this document.",
    )

    add_heading(doc, "7. TECHNOLOGIES USED", caps=True)
    add_body(doc, "SOFTWARE (from requirements.txt, with role in this project):")
    for pkg, desc in PIP_STACK:
        add_body(doc, f"• {pkg}: {desc}.")
    add_body(
        doc,
        "OTHER SOFTWARE USED IN SOURCE (not always duplicated in requirements as separate lines): Python standard library "
        "modules sqlite3, hashlib, secrets, smtplib, email, pathlib — for authentication storage and optional mail-based feedback.",
    )
    add_body(doc, "MODELS (TensorFlow Keras): " + ", ".join(MODEL_NAMES) + ".")
    add_body(
        doc,
        "DATASETS REFERENCED IN CODE: PH2 (diagnostic/gender/age plots) and ISIC2016 ground-truth CSV "
        "(ISBI2016_ISIC_Part3_Training_GroundTruth.csv) for label distribution plots when available.",
    )
    add_heading(doc, "HARDWARE", caps=True)
    add_body(
        doc,
        "Processor: multi-core CPU suitable for TensorFlow CPU builds. RAM: 8 GB or more. Storage: SSD recommended for faster "
        "weight loading.",
    )

    add_heading(doc, "8. SYSTEM STUDY AND TESTING", caps=True)
    add_heading(doc, "8.1 FEASIBILITY STUDY", caps=True)
    add_body(
        doc,
        "Economic feasibility: core libraries are open source. Technical feasibility: requirements.txt pins a workable TensorFlow "
        "CPU stack. Social feasibility: disclaimers and educational framing reduce misuse risk in demos.",
    )
    add_heading(doc, "8.2 TYPES OF TESTING", caps=True)
    add_body(
        doc,
        "8.2.1 Unit testing: password hashing, input validation, small preprocessing helpers.\n"
        "8.2.2 Integration testing: login flow with SQLite, navigation between Streamlit pages, model load then predict.\n"
        "8.2.3 Acceptance testing: sample images produce outputs when weights exist.\n"
        "8.2.4 Functional testing: each sidebar destination renders; invalid files are rejected.\n"
        "8.2.5 White box testing: missing weight files trigger handled errors.\n"
        "8.2.6 Black box testing: external tester uses only the browser, without reading code.",
    )
    add_heading(doc, "8.3 TEST CASES", caps=True)
    add_body(
        doc,
        "TC01 Valid registration — new email — account row created.\n"
        "TC02 Duplicate email — registration rejected with message.\n"
        "TC03 Valid login — session authenticated.\n"
        "TC04 Wrong password — login denied.\n"
        "TC05 Valid JPG upload — preview succeeds.\n"
        "TC06 Non-image upload — user sees guidance.\n"
        "TC07 Each model selection — preprocessing matches that backbone.\n"
        "TC08 Missing .keras/.h5 file — clear error, app stable.\n"
        "TC09 Visualization page — CSV present renders charts; missing CSV shows handled failure.\n"
        "TC10 Feedback form — submits without crashing (email depends on configuration).",
    )

    add_heading(doc, "9. CONCLUSION", caps=True)
    add_body(
        doc,
        "We delivered a Melanoma Detection Web Application that matches the real codebase: Streamlit plus TensorFlow models "
        f"({', '.join(MODEL_NAMES)}), SQLite authentication, and the visualization stack from requirements.txt. The workflow and "
        "architecture figures document how a user moves from login to prediction, and how software layers connect.",
    )

    add_heading(doc, "10. FUTURE ENHANCEMENT", caps=True)
    add_body(
        doc,
        "• Cloud hosting with managed secrets for email.\n"
        "• Automated tests in CI for model load smoke checks.\n"
        "• Additional explainability overlays where saliency weights exist.\n"
        "• Packaging with Docker for one-command demos.",
    )

    add_heading(doc, "11. REFERENCES", caps=True)
    add_body(
        doc,
        "[1] Esteva, A., et al. (2017). Dermatologist-level classification of skin cancer with deep neural networks. Nature.\n"
        "[2] Codella, N., et al. (2018). Skin lesion analysis toward melanoma detection challenge materials.\n"
        "[3] Tschandl, P., et al. (2018). HAM10000 dataset, Scientific Data.\n"
        "[4] Streamlit documentation — https://docs.streamlit.io/\n"
        "[5] TensorFlow documentation — https://www.tensorflow.org/\n"
        "[6] ISIC Archive — https://www.isic-archive.com/\n"
        "[7] Project requirements.txt (package versions as pinned).",
    )

    add_heading(doc, "12. APPENDIX", caps=True)
    add_body(
        doc,
        "Sample dependency list (exact file in repository):\n"
        + "\n".join(line for line in Path(OUT_DIR / "requirements.txt").read_text(encoding="utf-8").strip().splitlines()),
    )

    add_heading(doc, "13. BIO DATA", caps=True)
    add_body(doc, "[Attach standard biodata sheets as required by the department.]")


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    flow = FIG_DIR / "figure_3_5_workflow.png"
    arch = FIG_DIR / "figure_3_6_architecture.png"
    build_flowchart_png(flow)
    build_architecture_png(arch)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title_block(doc)
    add_certificate_declaration_ack(doc)
    add_index_and_lists(doc)
    add_main_chapters(doc, flow, arch)

    doc.save(DOCX_PATH)
    print(f"Wrote: {DOCX_PATH}")


if __name__ == "__main__":
    main()
